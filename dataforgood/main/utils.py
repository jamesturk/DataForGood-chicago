import os
import uuid
import folium
import geopandas as gpd
import numpy as np
import pandas as pd
import branca.colormap as cm

from django import forms
from dataforgood.settings import BASE_DIR
from django.db.models import Avg, Sum
from docx import Document
from langchain.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_openai import ChatOpenAI

from .models import (
    ContractRent_Main,
    ContractRent_Sub,
    Disability_Main,
    Disability_Sub,
    Enrollment_Main,
    Enrollment_Sub,
    HouseholdType_Main,
    HouseholdType_Sub,
    Insurance_Main,
    Insurance_Sub,
    MeanIncome_Main,
    MeanIncome_Sub,
    MedianAge_Main,
    MedianAge_Sub,
    MedianEarning_Main,
    MedianEarning_Sub,
    MedianIncome_Main,
    MedianIncome_Sub,
    Races_Main,
    Races_Sub,
    TractZipCode,
)


communityshape_path = os.path.join(BASE_DIR, "main/communityarea")
zipcodeshape_path = os.path.join(BASE_DIR, "main/zipcode")
censusshape_path = os.path.join(BASE_DIR, "main/censustracts")
html_path = os.path.join(BASE_DIR, "main/templates/maps")
docs_path = os.path.join(BASE_DIR, "main/templates/memos")

# HELPER FUNCTIONS FOR VIEWS.PY #

# Global variables that map form indicator options to specific models to conduct querying
MAIN_MODEL_MAPPING = {
    "Median Income in the Past 12 Months (inflation-adjusted)": MedianIncome_Main,
    "Mean Income in the Past 12 Months (inflation-adjusted)": MeanIncome_Main,
    "Aggregate Contract Rent": ContractRent_Main,
    "Total Number of Households": HouseholdType_Main,
    "Median Earnings in the Past 12 Months": MedianEarning_Main,
    "Population 3 years and over enrolled in school": Enrollment_Main,
    "Total Population With Disability": Disability_Main,
    "Insurance Coverage: Total Population": Insurance_Main,
    "Total Population and Race Group": Races_Main,
    "Median Age": MedianAge_Main,
}

SUB_MODEL_MAPPING = {
    "Median Income in the Past 12 Months (inflation-adjusted)": MedianIncome_Sub,
    "Mean Income in the Past 12 Months (inflation-adjusted)": MeanIncome_Sub,
    "Aggregate Contract Rent": ContractRent_Sub,
    "Total Number of Households": HouseholdType_Sub,
    "Median Earnings in the Past 12 Months": MedianEarning_Sub,
    "Population 3 years and over enrolled in school": Enrollment_Sub,
    "Total Population With Disability": Disability_Sub,
    "Insurance Coverage: Total Population": Insurance_Sub,
    "Total Population and Race Group": Races_Sub,
    "Median Age": MedianAge_Sub,
}

INDICATOR_UNIT_MAPPING = {
    "Median Income in the Past 12 Months (inflation-adjusted)": "US Dollars",
    "Mean Income in the Past 12 Months (inflation-adjusted)": "US Dollars",
    "Aggregate Contract Rent": "US Dollars",
    "Total Number of Households": "Number of Households",
    "Median Earnings in the Past 12 Months": "US Dollars",
    "Population 3 years and over enrolled in school": "Number of People",
    "Total Population With Disability": "Number of People",
    "Insurance Coverage: Total Population": "Number of People",
    "Total Population and Race Group": "Number of People",
    "Median Age": "Years",
}

AGGERGATE_OPERATORS = {
    "Median Income in the Past 12 Months (inflation-adjusted)": "Average",
    "Mean Income in the Past 12 Months (inflation-adjusted)": "Average",
    "Aggregate Contract Rent": "Average",
    "Total Number of Households": "Total",
    "Median Earnings in the Past 12 Months": "Average",
    "Population 3 years and over enrolled in school": "Total",
    "Total Population With Disability": "Total",
    "Insurance Coverage: Total Population": "Total",
    "Total Population and Race Group": "Total",
    "Median Age": "Average",
}

SUBGROUP_NAMES = {
    "mean_white": "White",
    "mean_black": "Black",
    "mean_hawai": "Native Hawaiian and Other Pacific Islander",
    "mean_ind_ala": "American Indian and Alaska Native",
    "mean_asia": "Asian",
    "mean_other": "Some Other Race",
    "median_white": "White",
    "median_hawai": "Native Hawaiian and Other Pacific Islander",
    "median_asia": "Asian",
    "median_ind_ala": "American Indian and Alaska Native",
    "median_other": "Some Other Race",
    "median_black": "Black",
    "less_high": "Less Than High School Graduate",
    "grad": "Graduate or Professional Degree",
    "high": "High School Graduate",
    "bachelor": "Bachelor's Degree",
    "college": "Some College and Associate's Degree",
    "kind_12": "Kindergarten to 12th Grade",
    "grad": "Graduate or Professional School",
    "college": "College, Undergraduate",
    "nursery": "Nursery School, Preschool",
    "ind_living": "With an Independent Living Difficulty",
    "ambulatory": "With an Ambulatory Difficulty",
    "self_care": "With a Self-Care Difficulty",
    "vision": "With a Vision Difficulty",
    "hearing": "With a Hearing Difficulty",
    "cognitive": "With a Cognitive Difficulty",
    "insured": "Insured",
    "uninsured": "Uninsured",
    "lower_rent": "Lower Contract Rent Quartile",
    "upper_rent": "Upper Contract Rent Quartile",
    "median_rent": "Median Contract Rent",
    "household_family": "Total Number of Family Household",
    "household_nonfamily": "Total Number of Non-family Household",
    "pop_hawai": "Native Hawaiian and  Other Pacific Islander",
    "pop_asia": "Asian",
    "pop_black": "Black",
    "pop_white": "White",
    "pop_two": "Two or More Races",
    "pop_other": "Some Other Race",
    "pop_ind_ala": "American Indian and Alaska Native",
    "male_age": "Male",
    "female_age": "Female",
}

# Source: Modified class, context, and instructions from
# https://github.com/abejburton/census_llm


class WriteMemo:
    """
    This class uses GPT 3.5 Turbo to explain the selected indicator data.
    """

    def __init__(self, indicator, geo_level, dictionary, describe, open_ai_key):
        self.indicator = indicator
        self.geo_level = geo_level
        self.dictionary = dictionary
        self.describe = describe

        template = """
        Role: You are a data analyst for a small nonprofit in Chicago.

        Context:
        Your goal is to write a memo about the data from  the American
        Community Survey's (ACS) 5-year estimate data, which are "period"
        estimates that represent data collected over a period of time.
        The primary advantage of using multiyear estimates is the increased
        statistical reliability of the data for less populated areas and small
        population subgroups. The data is focused on a selected indicator,
        a selected geographic level, and selected years, all of which are
        chosen by the user.

        You will receive 4 pieces of information to form a truthful analysis
        and write a professional memo that can be interpreted by a
        nontechnical audience.

        The first piece of information is a string value of the selected
        indicator that the analysis will be focused on.

        The second piece of information is the selected geographic level.
        There are three geographic levels, Census Tracts, Zip codes, and
        Community Areas, and the user only picks one. The data is only focused
        on the geographic levels in Chicago so do not mention
        other cities, states, or countries.

        The third piece of information is the python output of dictionary of
        a pandas dataframe of selected data from the American Community Survey.
        The dictionary's keys are the dataframe's column names and the
        dictionary's values are the dataframe's rows of values which correspond
        to the columns. The first column will be the geographic level of data
        that was selected by the user. The following columns will be the years
        selected by the user. The years will range from 2017 to 2022. The
        values in the year columns reflect the values of indicator selected
        by the user. The values in the first column reflect the selected
        census tracts, community areas, or zip codes.

        The fourth piece of information will be the results of calling the
        describe() method on the dataframe that is selected by the user (this
        is not across all the geographic areas of Chicago).
        Please do not conduct analysis on the first column of the dataframe.
        Please provide a description of the dataframe, including the mean,
        minimum, and maximum values in a year for the indicator.
        For each geographic area in each year in the
        dictionary please list what quartile it is in and be sure to mention
        which geographic area is associated with the value. If the geographic
        are in the specified year is the minimum or maximum value in the
        selected dataframe, point it out.

        If there are multiple years in the dataframe, point out if the value
        for the indicator for a geographic area is increasing or decreasing.
        Lastly, please provide insight about the selected indicator and
        what factors the nonprofit may want to investigate further that are
        related to the selected indicator and how a nonprofit could help
        improve the indicator values for geographic areas that have the lowest
        values.

        Instructions:
        Explain that you did analysis on an indicator (mention indicator name)
        using data from the American Community Survey's (ACS)
        5-year estimate data for certain geographic areas in Chicago (list
        out the selected geographic areas) in certain years (list out the
        selected years) and will be explaining the output of
        the analysis.

        Read the four pieces of information provided and explain any key
        findings you identify. Base your analysis solely on the details
        given in this prompt. Do not make assumptions; provide a
        clear and truthful analysis based on the provided information.
        -----

        Information: {information}
        """
        prompt = ChatPromptTemplate.from_template(template)
        model = ChatOpenAI(
            model="gpt-3.5-turbo", temperature=0, api_key=open_ai_key
        )
        self.response = prompt | model | StrOutputParser()

    def invoke(self):
        return self.response.invoke(
            {
                "information": [
                    self.indicator,
                    self.geo_level,
                    self.dictionary,
                    self.describe,
                ]
            }
        )


def save_memo(indicator, geo_level, memo):
    """
    This function saves the memo outputted by ChatGPT to a word file.

    Inputs:
        - indicator (str): name of indicator selected by the user
        - geo_level (str): geographic level of data selected  by the user
        - memo (str): string of ChatGPT's response

    Returns:
        - path (str): string of the path where the memo is saved
    """
    document = Document()

    # Font style
    style = document.styles["Normal"]
    style.font.name = "Calibri"

    # Adding Memo title and body
    document.add_heading(
        "Analysis of {} in Selected Chicago {}".format(indicator, geo_level), 0
    )
    document.add_paragraph(memo)

    os.makedirs(docs_path, exist_ok=True)
    path = docs_path + "/memo_{}.docx".format(uuid.uuid4())

    document.save(path)

    return path


class MainTable:
    """
    Class object to represent a main data table on the webapp.
    """

    def __init__(self, geographic_level, geographic_units, indicator, 
                 model, periods):
        """
        Inputs:
            geographic_level (str): geographic level selected by the user
                in the form
            geographic_units (list of str or int): geographic unit(s) corresponding
                to the geographic level selected by the user in the form
            indicator (str): name of indicator selected by the user in the form
            model (Django model object): corresponding model for the indicator
            preiods (list of str): periods(s) selected by the user

        Returns: None
        """
        # Query Variables
        self.geographic_level = geographic_level
        self.geographic_units = self.convert_list_to_tuple(geographic_units)
        self.indicator = indicator
        self.model = model
        self.periods = periods
        self.aggregate_operation = AGGERGATE_OPERATORS[self.indicator]
        self.annotation_cls = self.get_annotation_cls()
        self.years = self.convert_periods_to_years()

        # Main Table Output Variables
        self.table_title = self.create_table_title()
        self.headers = [self.geographic_level] + periods
        self.num_cols = len(periods)
        self.rows = self.create_rows()
        self.table = {"headers": self.headers, "rows": self.rows}

    def convert_list_to_tuple(self, query_lst):
        """
        Converts lists of variables (e.g. years or tracts) into tuples to
        run model query for table creation.
        Examples:
            - [2018, 2019] -> (2018, 2019)
            - [2018] -> (2018,)

        Inputs:
            query_lst (list): list of query variable(s)

        Returns:
            query_tup (tuple): tuple of query variable(s)
        """
        if len(query_lst) == 1:
            query_tup = (query_lst[0],)
        else:
            query_tup = sorted(tuple(query_lst))

        return query_tup

    def convert_periods_to_years(self):
        """
        Takes in the 5-year estimate period(s) selected by the user and retains
        only the ending year of each period selected to conduct query
        An example:
            ['2011-2015', '2016-2020'] -> [2015, 2020]
            ['2011-2015'] -> [2015]

        Inputs: None

        Returns:
            years (list of int): the ending year for each 5-year estimate period
                selected by the user
        """
        years = []
        for p in self.periods:
            # Only keeps the last four numbers in the period string
            years.append(int(p[5:]))

        years = self.convert_list_to_tuple(years)

        return years

    def create_table_title(self):
        """
        Creates a title for the main table based on indicator and year(s) selected
        by the user.

        Inputs: None

        Returns:
            indicator_formatted (str): readable indicator format and 5-year periods
                selected by the user
        """
        indicator_formatted = self.indicator.replace("_", " ").title() + " for "

        for idx, period in enumerate(self.periods):
            if idx == len(self.periods) - 1:
                indicator_formatted += period
            else:
                indicator_formatted += period + ", "

        return indicator_formatted

    def get_annotation_cls(self):
        """
        Determines the type of annotation (Avg or Sum) to use when conducting
        aggregate query operations

        Inputs: None

        Returns (Django query variable): 
            Avg (averages groupby values) OR Sum (totals grouby values)
        """
        if self.aggregate_operation == "Average":
            return Avg
        elif self.aggregate_operation == "Total":
            return Sum

    def conduct_query_city_level(self):
        """
        Conducts a Django model query to obtain the city-level average or sum
        for a given indicator for one or more periods.

        Inputs: None

        Returns:
            results (Django Queryset): list of query result instances
        """
        results = (
            self.model.objects.values("year")
            .filter(year__in=self.years)
            .annotate(agg_val=self.annotation_cls("value"))
            .order_by("year")
        )

        return results

    def conduct_query_zipcode_level(self, one_zipcode):
        """
        Conducts a Django model query to obtain the zipcode-level average or sum
        for a given indicator for one or more periods.

        Inputs:
            one_zipcode (int): a zipcode number selected by the user

        Returns:
            results (Django Queryset): list of query result instances
        """
        tracts_in_zipcode = list(
            TractZipCode.objects.filter(zip_code=one_zipcode)
            .values_list("tract_id")
            .distinct()
        )

        results = (
                self.model.objects.values("year")
                .filter(
                    census_tract_id__in=tracts_in_zipcode, year__in=self.years
                )
                .annotate(agg_val=self.annotation_cls("value"))
                .order_by("year")
            )

        return results

    def conduct_query_tract_level(self, one_tract):
        """
        Conducts a Django model query to obtain the tract-level value for a
        given indicator for one or more periods.

        Inputs:
            one_tract (int): a tract number selected by the user

        Returns:
            results (Django Queryset): list of query result instances
        """
        results = self.model.objects.filter(
            census_tract_id=one_tract, year__in=self.years
        )

        return results

    def conduct_query_community_level(self, one_community):
        """
        Conducts a Django model query to obtain the community-level value for a
        given indicator for one or more periods.

        Inputs:
            one_community (str): a community selected by the user

        Returns:
            results (Django Queryset): list of query result instances
        """
        results = (
                self.model.objects.values("census_tract_id__community", "year")
                .filter(
                    census_tract_id__community=one_community,
                    year__in=self.years,
                )
                .annotate(agg_val=self.annotation_cls("value"))
                .order_by("year")
            )

        return results

    def create_one_row(self, results, unit):
        """
        Creates one row of the table for a given geographic unit.

        Inputs:
            results (Django Queryset): list of query result instances
            unit (str ot int): a single geographic unit

        Returns
            row (list): query result for one geographic unit across multiple
                periods
        """
        # Create a row of "NA" strongs corresponding to the number of
        # table columns (i.e. number of periods selected by the user)
        row = ["NA"] * self.num_cols

        if self.geographic_level == "Tract":
            for r in results:
                result_year = r.year
                idx = self.years.index(result_year)
                if r.value is None:
                    continue
                else:
                    row[idx] = round(float(r.value), 2)

        else:
            for r in results:
                result_year = r["year"]
                idx = self.years.index(result_year)
                if r["agg_val"] is None:
                    continue
                else:
                    row[idx] = round(float(r["agg_val"]), 2)
        
        if self.geographic_level == "Community":
            row = [unit.title()] + row
        else:
            row = [unit] + row

        return row

    def create_rows(self):
        """
        Generates a dictionary to be used a context variables in the html file to
        create a table on the webapp (for the main indicator/overall group).

        Inputs: None

        Returns:
            rows (list of lists): Each list represents one row/one geographic
                unit (city, tract, zipcode community)
        """
        rows = []

        if self.geographic_level == "City of Chicago":
            results = self.conduct_query_city_level()
            row = self.create_one_row(
                results, "City " + self.aggregate_operation
            )

            rows.append(row)

        for unit in self.geographic_units:
            if self.geographic_level == "Tract":
                results = self.conduct_query_tract_level(unit)
                row = self.create_one_row(results, unit)

            elif self.geographic_level == "Zipcode":
                results = self.conduct_query_zipcode_level(unit)
                row = self.create_one_row(results, unit)

            elif self.geographic_level == "Community":
                results = self.conduct_query_community_level(unit)
                row = self.create_one_row(results, unit)

            rows.append(row)

        return rows


class SubgroupTable:
    """
    Class object to represent a series of subgroup data tables on the webapp.
    """

    def __init__(self, geographic_level, geographic_units, indicator, 
                 model, periods):
        """
        Inputs:
            geographic_level (str): geographic level selected by the user
                in the form
            geographic_units (list of str or int): geographic unit(s) corresponding
                to the geographic level selected by the user in the form
            indicator (str): name of indicator selected by the user in the form
            model (Django model object): corresponding model for the indicator
            preiods (list of str): periods(s) selected by the user

        Returns: None
        """
        # Query Variables
        self.geographic_level = geographic_level
        self.geographic_units = self.convert_list_to_tuple(geographic_units)
        self.indicator = indicator
        self.model = model
        self.periods = periods
        self.aggregate_operation = AGGERGATE_OPERATORS[self.indicator]
        self.annotation_cls = self.get_annotation_cls()
        self.years = self.convert_periods_to_years()
        self.model_subgroups = self.get_model_subgroups()

        # Subgroup Table Output Variables
        self.num_cols = self.get_num_cols()
        self.subtable_headers = self.create_subtable_headers()
        self.many_subtables = self.create_all_years_tables()

    def get_num_cols(self):
        """
        Returns the number of geographic units in the query, one for each column

        Inputs: None
        Returns (int): Number of columns
        """
        if self.geographic_level == "City of Chicago":
            return 1
        else:
            return len(self.geographic_units)
    
    def create_subtable_headers(self):
        """
        Creates the header row of the table.

        Inputs: None

        Returns:
            headers (list of str): table header list
        """
        headers = []
        if self.geographic_level == "City of Chicago":
            if self.aggregate_operation == "Average":
                headers = ["City of Chicago", "City Average"]
            elif self.aggregate_operation == "Total":
                headers = ["City of Chicago", "City Total"]
        
        elif self.geographic_level == "Community":
            headers.append(self.geographic_level)
            for community in self.geographic_units:
                headers.append(community.title())

        else:
            headers.append(self.geographic_level)
            headers.extend(self.geographic_units)

        return headers

    def convert_list_to_tuple(self, query_lst):
        """
        Converts lists of variables (e.g. years or tracts) into tuples to
        run model query for table creation.
        Examples:
            - [2018, 2019] -> (2018, 2019)
            - [2018] -> (2018,)

        Inputs:
            query_lst (list): list of query variable(s)

        Returns:
            query_tup (tuple): tuple of query variable(s)
        """
        if len(query_lst) == 1:
            query_tup = (query_lst[0],)
        else:
            query_tup = tuple(query_lst)

        return query_tup

    def convert_periods_to_years(self):
        """
        Takes in the 5-year estimate period(s) selected by the user and retains
        only the ending year of each period selected to conduct query
        An example:
            ['2011-2015', '2016-2020'] -> [2015, 2020]
            ['2011-2015'] -> [2015]

        Inputs: None

        Returns:
            years (list of int): the ending year for each 5-year estimate period
                selected by the user
        """
        years = []
        for p in self.periods:
            # Only keeps the last four numbers in the period string
            years.append(int(p[5:]))

        return self.convert_list_to_tuple(years)

    def get_model_subgroups(self):
        """
        Generates a list of unique subgroups for selected geographic units
        for a given indicator.

        Inputs: None

        Returns:
            subgroups_lst (list of str): a list of unique subgroups
        """
        subgroups = self.model.objects.values_list(
            "sub_group_indicator_name"
        ).distinct()

        subgroups_lst = []
        for subgroup in subgroups:
            subgroups_lst.append(subgroup[0])

        return subgroups_lst

    def get_annotation_cls(self):
        """
        Determines the type of annotation (Avg or Sum) to use when conducting
        aggregate query operations

        Inputs: None

        Returns (Django query variable): 
            Avg (averages groupby values) OR Sum (totals grouby values)
        """
        if self.aggregate_operation == "Average":
            return Avg
        elif self.aggregate_operation == "Total":
            return Sum
    
    def conduct_query_city_level(self, one_year):
        """
        Conducts a Django model query to obtain the city-level average or sum
        for the various subgroups of an indicator, for a given year.

        Inputs:
            one_year (int): a year/period selected by the user

        Returns:
            results (Django Queryset): list of query result instances
        """
        results = (
                self.model.objects.values("sub_group_indicator_name")
                .filter(year=one_year)
                .annotate(agg_val=self.annotation_cls("value"))
                .order_by("sub_group_indicator_name")
            )

        return results

    def conduct_query_zipcode_level(self, zipcode, one_year):
        """
        Conducts a Django model query to obtain the zipcode-level average or sum
        for the various subgroups of an indicator, for a given year.

        Inputs:
            one_zipcode (int): a single zipcode selected by the user
            one_year (int): a year/period selected by the user

        Returns:
            results (Django Queryset): list of query result instances
        """
        tracts_in_zipcode = list(
            TractZipCode.objects.filter(zip_code=zipcode)
            .values_list("tract_id")
            .distinct()
        )

        results = (
                self.model.objects.filter(
                    census_tract_id__in=tracts_in_zipcode, year=one_year
                )
                .values("sub_group_indicator_name")
                .annotate(agg_val=self.annotation_cls("value"))
                .order_by("sub_group_indicator_name")
            )

        return results

    def conduct_query_community_level(self, one_year):
        """
        Conducts a Django model query to obtain the community-level average or sum
        for the various subgroups of an indicator, for a given year.

        Inputs:
            one_year (int): a year/period selected by the user

        Returns:
            results (Django Queryset): list of query result instances
        """
        results = (
                self.model.objects.filter(
                    census_tract_id__community__in=self.geographic_units,
                    year=one_year,
                )
                .values(
                    "census_tract_id__community", "sub_group_indicator_name"
                )
                .annotate(agg_val=self.annotation_cls("value"))
                .order_by(
                    "sub_group_indicator_name", "census_tract_id__community"
                )
            )

        return results

    def conduct_query_tract_level(self, one_year):
        """
        Conducts a Django model query to obtain the tract-level average or sum
        for the various subgroups of an indicator, for a given year.

        Inputs:
            one_year (int): a year/period selected by the user

        Returns:
            results (Django Queryset): list of query result instances
        """
        results = (
            self.model.objects.filter(
                census_tract_id__in=self.geographic_units, year=one_year
            )
            .values("census_tract_id", "sub_group_indicator_name")
            .annotate(agg_val=Avg("value"))
            .order_by("sub_group_indicator_name", "census_tract_id")
        )

        return results

    def create_rows(self, results, rows):
        """
        Creates multiple rows, one row for each subgroup.

        Inputs:
            subgroup_lst (list of str): list of subgroup names
            rows (list of lists): a single list to store each row as a list
            results (Django queryset): list of query result instances

        Returns:
            rows (list of lists): a single list with each subgroup stored as a
                list
        """     
        for subgroup in self.model_subgroups:
            row = ["NA"] * self.num_cols
    
            if self.geographic_level == "City of Chicago":
                for r in results.filter(sub_group_indicator_name=subgroup):
                    if r["agg_val"] is None:
                        continue
                    else:
                        row[0] = round(float(r["agg_val"]), 2)

            else:
                for r in results.filter(sub_group_indicator_name=subgroup):

                    if self.geographic_level == "Tract":
                        result_geographic_area = str(r["census_tract_id"])
                    
                    elif self.geographic_level == "Community":
                        result_geographic_area = str(r["census_tract_id__community"])
                    
                    idx = self.geographic_units.index(result_geographic_area)
                    
                    if r["agg_val"] is None:
                        continue
                    else:
                        row[idx] = round(float(r["agg_val"]), 2)
            
            row = [SUBGROUP_NAMES[subgroup]] + row
            rows.append(row)

        return rows

    def create_all_years_tables(self):
        """
        Creates multiple subgroup tables, one table for each year.

        Inputs: None

        Returns:
            all_years_tables (dict): A nested dictionary, each interior dictionary
                representing one period selected by the user
                - Key (str): periods selected by the user
                - Value (dict): one subgroup table for a given period
        """
        all_years_tables = {period: None for period in self.periods}
        for period_str, one_year in zip(self.periods, self.years):
            one_year_table = self.create_one_year_table(one_year)
            all_years_tables[period_str] = one_year_table

        return all_years_tables

    def create_one_year_table(self, one_year):
        """
        Generates one subgroup table, for a given year.

        Inputs:
            one_year (int): a single period/year

        Returns (dict): a dictionary representing the headers and rows of a table
            - 'headers': header row of table (list of str)
            - 'rows': multiple rows for each geographic unit (list of lists of
                str)
        """
        rows = []

        if self.geographic_level == "City of Chicago":
            results = self.conduct_query_city_level(one_year)
            rows = self.create_rows(results, rows)

        if self.geographic_level == "Tract":
            results = self.conduct_query_tract_level(one_year)
            rows = self.create_rows(results, rows)

        if self.geographic_level == "Zipcode":
            # Create a dictionary to save values for each subgroup
            subgroup_dct = {
                subgroup: ["NA"] * self.num_cols
                for subgroup in self.model_subgroups
            }

            # Conduct querying for each zipcode
            for idx, zipcode in enumerate(self.geographic_units):
                results = self.conduct_query_zipcode_level(zipcode, one_year)
                for r in results:
                    if r["agg_val"] is None:
                        continue
                    else:
                        subgroup_indicator_name = r["sub_group_indicator_name"]
                        subgroup_dct[subgroup_indicator_name][idx] = round(
                            r["agg_val"], 2
                        )

            # Convert dictionary values to list of lists for table
            rows = [
                [SUBGROUP_NAMES[subgroup]] + row
                for subgroup, row in subgroup_dct.items()
            ]

        if self.geographic_level == "Community":
            results = self.conduct_query_community_level(one_year)
            rows = self.create_rows(results, rows)

        return {"headers": self.subtable_headers, "rows": sorted(rows)}


def generate_heatmaps(geograpahic_level, indicator, field, years):
    """
    Generates heat maps for the specified geographic level and indicator.

    Args:
        geograpahic_level (str): The geographic level (Community, Zipcode, Tract).
        indicator (str): The selected indicator.
        field (dict): The field data containing headers and rows.
        years (list): The list of years selected.

    Returns:
        heatmap_data: Pandas DataFrame that is used to generate the heatmap.
        heatmap_info: A list of dictionaries containing heatmap information.
    """
    heatmap_data = pd.DataFrame(field["rows"], columns=field["headers"])
    heatmap_info = []
    os.makedirs(html_path, exist_ok=True)

    for column in heatmap_data.columns[1:]:
        heatmap_data[column] = heatmap_data[column].apply(
            lambda x: float(x) if x != "NA" else np.nan
        )

    if geograpahic_level != "City of Chicago":
        if geograpahic_level == "Community":
            heatmap_data.iloc[:, 0] = heatmap_data.iloc[:, 0].str.upper()
            geo = gpd.read_file(communityshape_path)
            data = pd.merge(geo, heatmap_data, left_on="community", 
                            right_on="Community")

        elif geograpahic_level == "Zipcode":
            geo = gpd.read_file(zipcodeshape_path)
            data = pd.merge(geo, heatmap_data, left_on="zip", 
                            right_on="Zipcode")

        elif geograpahic_level == "Tract":
            geo = gpd.read_file(censusshape_path)
            data = pd.merge(geo, heatmap_data, left_on="tractce10", 
                            right_on="Tract")

        for year in years:
            year_dic = {}
            heatmap, title = create_heatmap(data, geograpahic_level, indicator, 
                                          year)
            
            name = "heatmap_{}".format(uuid.uuid4())
            map_file_path = "{}/{}.html".format(html_path, name)
            heatmap.save(map_file_path)
            year_dic["title"] = title
            year_dic["path"] = f"maps/{name}.html"
            year_dic["year"] = year
            heatmap_info.append(year_dic)

    return heatmap_data, heatmap_info

def delete_legend(Choropleth: folium.Choropleth):
  """
  This function removes the legend from the choropleth. This function is based
  from https://github.com/python-visualization/folium/issues/1052
  
  Args:
    choropleth: Folium choropleth object
    
  Returns:
    Chloropeth without legend
  """
  legend_list = []
  for c in Choropleth._children:
    if c.startswith('color_map'):
      legend_list.append(c)
  for item in legend_list:
    Choropleth._children.pop(item)
  return Choropleth

def create_heatmap(data, geograpahic_level, indicator, year):
    """
    Creates a heatmap using Folium.

    Args:
        data (GeoDataFrame): The merged geographic and field data.
        geograpahic_level (str): The geographic level (Community, Zipcode, Tract).
        indicator (str): The selected indicator.
        year (str): The selected year.

    Returns:
        heatmap: The generated folium heatmap
        title: Title of folium heatmap
    """
    # Adding base layer
    heatmap = folium.Map(
        location=[41.8781, -87.6298], zoom_start=10, tiles=None
    )
    folium.TileLayer(
        "CartoDB positron", name="Light Map", control=False
    ).add_to(heatmap)

    # Creating heatmap layer
    units = INDICATOR_UNIT_MAPPING[indicator]

    delete_legend(folium.Choropleth(
        geo_data=data,
        name="Choropleth",
        data=data,
        columns=[geograpahic_level, year],
        key_on="feature.properties.{}".format(geograpahic_level),
        fill_color="Blues",
        fill_opacity=1,
        line_opacity=0.2,
        bins=3,
        legend_name=units,
        smooth_factor=0,
        nan_fill_color="grey",
        nan_fill_opacity=0.4,
    )).add_to(heatmap)

    def style_function(x):
        return {
            "fillColor": "#ffffff",
            "color": "#000000",
            "fillOpacity": 0.1,
            "weight": 0.1,
        }

    def highlight_function(x):
        return {
            "fillColor": "#000000",
            "color": "#000000",
            "fillOpacity": 0.5,
            "weight": 0.1,
        }

    # Adding tooltips to map
    feat = folium.features.GeoJson(
        data,
        style_function=style_function,
        control=False,
        highlight_function=highlight_function,
        tooltip=folium.features.GeoJsonTooltip(
            fields=[geograpahic_level, year],
            aliases=[geograpahic_level, units],
            style=(
                "background-color: white; color: #333333; font-family: arial; font-size: 12px; padding: 10px;"
            ),
        ),
    )
    heatmap.add_child(feat)
    heatmap.keep_in_front(feat)
    folium.LayerControl().add_to(heatmap)
    if geograpahic_level[-1:] == "y":
        title = "Heat Map of {} by Selected {}ies in {}".format(
            indicator, geograpahic_level[:-1], year
        )
    else:
        title = "Heat Map of {} by Selected {}s in {}".format(
            indicator, geograpahic_level, year
        )
    return heatmap, title


# HELPER FUNCTIONS FOR FORMS.PY #


def get_choices(model, col):
    """
    Generates a list of unique select options based on data in a column of a
        selected model.

    Inputs:
        model (model object): model based on the category selected by the user
        col (str): name of column in the model

    Returns:
        choices (list of tuples): corresponding choices as a list to be used in
            the user search form
    """
    choices = []
    unique_values = list(model.objects.values(col).distinct())
    for item in unique_values:
        choices.append((item[col], item[col]))

    return sorted(choices)


def create_multiple_choice_geo(
    model, field_name, widget_id, initial_val, class_name="form-option"
):
    """
    Create a dynamically configured MultipleChoiceField based on the specified model and field.

    :param model: Django model class from which to fetch choices.
    :param field_name: Name of the model field for which to create a choice field.
    :param widget_id: HTML ID to use for the field's widget.
    :param initial_val: initial or default value of the field
    :param class_name: CSS class name for the widget.
    :return: forms.MultipleChoiceField instance.
    """
    return forms.MultipleChoiceField(
        initial=initial_val,
        widget=forms.CheckboxSelectMultiple(
            attrs={"class": class_name, "id": widget_id},
        ),
        choices=get_choices(model, field_name),
    )


def create_multiple_choice_indicator(
    choices_lst, widget_id, class_name="form-option"
):
    """
    Create a dynamically configured MultipleChoiceField based on the specified model and field.

    :param model: Django model class from which to fetch choices.
    :param field_name: Name of the model field for which to create a choice field.
    :param widget_id: HTML ID to use for the field's widget.
    :param class_name: CSS class name for the widget.
    :return: forms.MultipleChoiceField instance.
    """
    return forms.ChoiceField(
        widget=forms.Select(attrs={"class": class_name, "id": widget_id}),
        choices=choices_lst,
    )

def prepare_chart_data(field):
    chart_data = {
        "categories": field["headers"][1:],  # Years
        "series": [],
    }
    for row in field["rows"]:
        chart_data["series"].append(
            {
                "name": row[0],  # Geographic unit
                "data": row[1:],  # Values for each year
            }
        )
    return chart_data


def prepare_subgroup_chart_data(multi_year_subtable_field):
    subgroup_chart_data = {}
    for year_value, subtable_data in multi_year_subtable_field.items():
        subgroup_chart_data[year_value] = {
            "categories": subtable_data["headers"][1:],  # Subgroup categories
            "series": [
                {
                    "name": subtable_data["headers"][0],
                    "data": [row[1:] for row in subtable_data["rows"]],  # Subgroup values
                }
            ],
        }
    return subgroup_chart_data